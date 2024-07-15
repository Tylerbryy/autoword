import json
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_apa_document(title, author, institution, course, instructor, due_date, abstract, keywords, content, references, is_professional=False):
    doc = docx.Document()
    
    # Set up the document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set font for entire document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Enable double spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 2.0

    # Title Page
    if is_professional:
        running_head = title.upper()[:50]  # Limit to 50 characters
        header = sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"Running head: {running_head}"
        header_para.style = doc.styles['Header']
    
    # Center all content vertically on the title page
    for _ in range(10):  # Add some empty paragraphs to push content down
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True

    for item in [author, institution, course, instructor, due_date]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(item)

    doc.add_page_break()

    # Abstract (only for professional papers or if specifically requested)
    if is_professional or abstract:
        abstract_para = doc.add_paragraph('Abstract')
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abstract_para.runs[0].bold = True
        
        doc.add_paragraph(abstract)
        
        # Keywords
        keywords_para = doc.add_paragraph('Keywords: ')
        keywords_para.add_run(', '.join(keywords)).italic = True

        doc.add_page_break()

    # Content
    content_paragraphs = content.split('\n')
    in_list = False
    list_item_prefixes = ['-', '•', '*']  # Common list item prefixes
    
    for i, paragraph in enumerate(content_paragraphs):
        p = doc.add_paragraph()
        if i == 0:  # Center the first paragraph (Introduction)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Detect and handle list items
        stripped_paragraph = paragraph.strip()
        if any(stripped_paragraph.startswith(prefix) for prefix in list_item_prefixes):
            if not in_list:
                in_list = True
                p.style = 'List Bullet'
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
            else:
                p.style = 'List Bullet'
            paragraph = stripped_paragraph[1:].strip()  # Remove the list item prefix
        elif stripped_paragraph and stripped_paragraph[0].isdigit() and '.' in stripped_paragraph[:3]:
            if not in_list:
                in_list = True
                p.style = 'List Number'
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
            else:
                p.style = 'List Number'
        else:
            if in_list:
                in_list = False
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.first_line_indent = Inches(0.5)
        
        p.add_run(paragraph)

    doc.add_page_break()

    # References
    references_heading = doc.add_paragraph("References")
    references_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    references_heading.runs[0].bold = True

    for ref in references:
        para = doc.add_paragraph(ref)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)

    # Add page numbers
    add_page_number(doc.sections[0])

    return doc

def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    run = paragraph.add_run()
    run._element.append(field)
    
    return paragraph

# Load data from JSON file
with open('essay_data.json', 'r') as f:
    data = json.load(f)

title = data['title']
author = data['author']
institution = data['institution']
course = data['course']
instructor = data['instructor']
due_date = data['due_date']
abstract = data['abstract']
keywords = data['keywords']
content = data['content']
references = data['references']

# Set is_professional to True for professional papers, False for student papers
is_professional = False

doc = create_apa_document(title, author, institution, course, instructor, due_date, abstract, keywords, content, references, is_professional)
doc.save('essay.docx')